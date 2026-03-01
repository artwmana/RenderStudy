from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Cm

from . import gost_format
from . import markdown_parser, renderer_docx

NON_NUMBERED_TITLES = {
    "СОДЕРЖАНИЕ",
    "РЕФЕРАТ",
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "ПРИЛОЖЕНИЯ",
}


def reformat_docx(input_path: str | Path, output_path: str | Path) -> None:
    """Normalize an existing DOCX to project GOST formatting.

    The first page (title page) is preserved as-is. Formatting is applied
    starting from the next page after the first explicit page break.
    """
    src = Path(input_path)
    out = Path(output_path)
    doc = DocxDocument(str(src))
    gost_format.apply_page_layout(doc)

    paragraphs = list(doc.paragraphs)
    body_start_idx = _find_body_start_index(paragraphs)

    for idx, paragraph in enumerate(paragraphs):
        if idx < body_start_idx:
            continue
        _format_paragraph(paragraph)

    # Remove visually empty spacer paragraphs in body part.
    for idx, paragraph in list(enumerate(list(doc.paragraphs))):
        if idx < body_start_idx:
            continue
        if _is_removable_blank(paragraph):
            _remove_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _format_table_paragraph(paragraph)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))


def rebuild_docx_via_markdown(
    input_path: str | Path,
    output_path: str | Path,
    extracted_md_path: str | Path | None = None,
    title_template_path: str | Path | None = None,
) -> None:
    """Rebuild DOCX as: external title template + body text converted through Markdown.

    The first page of the input DOCX is treated as a removable title page and excluded.
    """
    src = Path(input_path)
    out = Path(output_path)
    original = DocxDocument(str(src))
    paragraphs = _collect_body_paragraphs_for_extraction(original)
    split_idx = _find_title_split_index(paragraphs)
    template = _resolve_title_template_path(title_template_path)
    with tempfile.TemporaryDirectory(prefix="renderstudy_docx_") as tmp_dir:
        tmp_root = Path(tmp_dir)
        images_dir = tmp_root / "images"
        images_dir.mkdir(parents=True, exist_ok=True)
        md_text = _extract_body_markdown(paragraphs, split_idx + 1, images_dir=images_dir)

        if extracted_md_path is not None:
            Path(extracted_md_path).write_text(md_text, encoding="utf-8")

        document = markdown_parser.parse_markdown(md_text)
        renderer_docx.render_document(
            document,
            output_path=out,
            asset_root=tmp_root,
            template_path=template,
        )


def _format_paragraph(paragraph) -> None:
    text = paragraph.text.strip()
    is_blank = text == ""
    is_non_numbered_heading = text.upper() in NON_NUMBERED_TITLES and text != ""
    is_numbered_heading = _looks_like_numbered_heading(text)

    if is_non_numbered_heading:
        gost_format.apply_heading_format(paragraph, centered=True, with_indent=False)
    elif is_numbered_heading:
        gost_format.apply_heading_format(paragraph, centered=False, with_indent=True)
    else:
        gost_format.apply_body_paragraph_format(paragraph)
        if is_blank:
            paragraph.paragraph_format.first_line_indent = Cm(0)

    if is_blank and not paragraph.runs:
        run = paragraph.add_run(" ")
        gost_format.set_run_font(run)
    else:
        for run in paragraph.runs:
            gost_format.set_run_font(run, bold=bool(run.bold), italic=bool(run.italic))


def _format_table_paragraph(paragraph) -> None:
    gost_format.apply_body_paragraph_format(paragraph)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    if not paragraph.runs:
        run = paragraph.add_run(" ")
        gost_format.set_run_font(run)
    else:
        for run in paragraph.runs:
            gost_format.set_run_font(run, bold=bool(run.bold), italic=bool(run.italic))


def _looks_like_numbered_heading(text: str) -> bool:
    if not text:
        return False
    return re.match(r"^\d+(?:\.\d+)*\s+\S", text) is not None


def _find_title_split_index(paragraphs: list) -> int:
    for idx, paragraph in enumerate(paragraphs):
        if _contains_page_break(paragraph):
            return idx
    # Fallback without explicit page break:
    # find first likely body-start paragraph and cut everything before it.
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if _is_body_start_candidate(paragraph, text):
            return idx - 1
    # Last fallback: keep old behavior.
    return 0 if paragraphs else -1


def _extract_body_markdown(paragraphs: list, start_idx: int, images_dir: Path | None = None) -> str:
    body_paragraphs = _trim_old_title_tail(paragraphs[start_idx:])
    lines: list[str] = []
    code_buffer: list[str] = []
    in_list = False
    semicolon_list_mode = False
    image_counter = 1
    idx = 0

    def flush_code_buffer() -> None:
        if not code_buffer:
            return
        # Use tilde fence to avoid collisions with backticks present in source text.
        lines.append("~~~~")
        lines.extend(code_buffer)
        lines.append("~~~~")
        lines.append("")
        code_buffer.clear()

    while idx < len(body_paragraphs):
        paragraph = body_paragraphs[idx]
        image_refs: list[str] = []
        if images_dir is not None:
            image_refs, image_counter = _extract_paragraph_images(paragraph, images_dir, image_counter)
        text = _normalize_markdown_artifacts(paragraph.text)
        formula_text = _extract_formula_text(paragraph, text)
        list_kind = _detect_list_kind(paragraph)
        if image_refs:
            if in_list:
                lines.append("")
                in_list = False
            semicolon_list_mode = False
            flush_code_buffer()
            image_caption = _extract_figure_caption(text)
            skip_next_caption = False
            if image_caption is None:
                for look_ahead in range(idx + 1, min(len(body_paragraphs), idx + 6)):
                    next_text = body_paragraphs[look_ahead].text.strip()
                    if not next_text:
                        continue
                    image_caption = _extract_figure_caption(next_text)
                    if image_caption is not None:
                        skip_next_caption = True
                        break
                    # Stop scan at first non-empty non-caption paragraph.
                    break
            for image_ref in image_refs:
                if image_caption:
                    lines.append(f'![Иллюстрация]({image_ref} "{image_caption}")')
                else:
                    lines.append(f"![Иллюстрация]({image_ref})")
                lines.append("")
            if skip_next_caption:
                idx += 1
            if not text:
                idx += 1
                continue
            if image_caption is not None:
                # Caption was embedded in this paragraph; it is now attached to image markdown.
                idx += 1
                continue
        if not text:
            if in_list:
                lines.append("")
                in_list = False
            semicolon_list_mode = False
            if formula_text:
                flush_code_buffer()
                lines.append(f"$${formula_text}$$")
                lines.append("")
                continue
            flush_code_buffer()
            if lines and lines[-1] != "":
                lines.append("")
            idx += 1
            continue
        if _is_markdown_fence_line(text):
            # Fence markers from previously exported markdown should not leak into
            # extracted markdown; otherwise they may swallow following content.
            idx += 1
            continue
        if formula_text:
            if in_list:
                lines.append("")
                in_list = False
            semicolon_list_mode = False
            flush_code_buffer()
            lines.append(f"$${formula_text}$$")
            lines.append("")
            idx += 1
            continue
        heading_level = _detect_heading_level(paragraph, text)
        if heading_level is not None:
            if in_list:
                lines.append("")
                in_list = False
            semicolon_list_mode = False
            flush_code_buffer()
            lines.append(f"{'#' * heading_level} {text}")
            lines.append("")
            idx += 1
        else:
            formatted = _format_term_line(text)
            if _looks_like_code_text(formatted):
                if in_list:
                    lines.append("")
                    in_list = False
                semicolon_list_mode = False
                code_buffer.append(formatted)
                idx += 1
                continue
            # Rule priority: lines inferred from ":" + ";" must be bullets.
            if semicolon_list_mode and _looks_like_semicolon_item(formatted):
                flush_code_buffer()
                lines.append(f"- {formatted}")
                in_list = True
                # Dot usually marks the final item in such list.
                if formatted.strip().endswith("."):
                    semicolon_list_mode = False
                idx += 1
                continue
            if list_kind is not None:
                flush_code_buffer()
                marker = "1." if list_kind == "ordered" else "-"
                lines.append(f"{marker} {formatted}")
                in_list = True
                semicolon_list_mode = False
                idx += 1
                continue
            if in_list:
                lines.append("")
                in_list = False
            flush_code_buffer()
            if _should_join_with_previous(lines, formatted):
                target_idx = _last_non_empty_index(lines)
                if target_idx is not None:
                    lines[target_idx] = f"{lines[target_idx].rstrip()} {formatted.lstrip()}"
                    del lines[target_idx + 1 :]
            else:
                lines.append(formatted)
                lines.append("")
            semicolon_list_mode = formatted.endswith(":")
            idx += 1
    if in_list:
        lines.append("")
    flush_code_buffer()
    while lines and lines[-1] == "":
        lines.pop()
    return "\n".join(lines)


def _detect_heading_level(paragraph, text: str) -> int | None:
    clean = text.strip()
    if not clean:
        return None
    # Prefer explicit Word heading styles if present.
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    heading_match = re.match(r"Heading\s+([1-6])", style_name, flags=re.IGNORECASE)
    if heading_match:
        return int(heading_match.group(1))
    if clean.upper() in NON_NUMBERED_TITLES:
        return 1
    # Uppercase headings (caps) should map to markdown heading.
    if _looks_like_caps_heading(clean):
        return 1
    if _looks_like_numbered_heading(clean):
        num = clean.split(maxsplit=1)[0]
        level = num.count(".") + 1
        return min(max(level, 1), 6)
    return None


def _trim_old_title_tail(paragraphs: list) -> list:
    """Drop leading leftovers from old title page (usually before 'цель')."""
    if not paragraphs:
        return paragraphs

    first_non_blank = _first_non_blank_index(paragraphs)
    if first_non_blank is None:
        return []
    trimmed = paragraphs[first_non_blank:]

    # Try to cut to the block around "цель" if found early.
    goal_idx = _find_goal_index(trimmed)
    if goal_idx is not None:
        # Keep nearest heading before goal, otherwise start at goal line.
        heading_before = _find_last_heading_before(trimmed, goal_idx)
        start = heading_before if heading_before is not None else goal_idx
        return trimmed[start:]

    # Fallback: start from first heading-like line if present.
    first_heading = _find_first_body_heading_index(trimmed)
    if first_heading is not None:
        return trimmed[first_heading:]
    return trimmed


def _first_non_blank_index(paragraphs: list) -> int | None:
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() or _paragraph_has_image(paragraph):
            return idx
    return None


def _find_goal_index(paragraphs: list) -> int | None:
    for idx, paragraph in enumerate(paragraphs[:40]):
        if re.search(r"\bцель\b", paragraph.text, flags=re.IGNORECASE):
            return idx
    return None


def _find_first_heading_index(paragraphs: list) -> int | None:
    for idx, paragraph in enumerate(paragraphs):
        if _detect_heading_level(paragraph, paragraph.text) is not None:
            return idx
    return None


def _find_first_body_heading_index(paragraphs: list) -> int | None:
    for idx, paragraph in enumerate(paragraphs):
        if _is_heading_candidate(paragraph, paragraph.text):
            return idx
    return None


def _find_last_heading_before(paragraphs: list, idx_limit: int) -> int | None:
    for idx in range(idx_limit, -1, -1):
        paragraph = paragraphs[idx]
        if _is_heading_candidate(paragraph, paragraph.text):
            return idx
    return None


def _is_body_start_candidate(paragraph, text: str) -> bool:
    clean = text.strip()
    if not clean:
        return False
    low = clean.lower()
    if low.startswith("цель"):
        return True
    if clean.upper() in NON_NUMBERED_TITLES:
        return True
    if _looks_like_numbered_heading(clean):
        return True
    return _is_heading_candidate(paragraph, clean)


def _is_heading_candidate(paragraph, text: str) -> bool:
    clean = text.strip()
    if not clean:
        return False
    if clean.upper() in NON_NUMBERED_TITLES:
        return True
    if _looks_like_numbered_heading(clean):
        return True
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    heading_match = re.match(r"Heading\s+([1-6])", style_name, flags=re.IGNORECASE)
    if heading_match:
        return True
    return False


def _detect_list_kind(paragraph) -> str | None:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    lower_style = style_name.lower()
    if "bullet" in lower_style:
        return "unordered"
    if "number" in lower_style or "номер" in lower_style:
        return "ordered"
    p_pr = paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        # Without reading numbering.xml fully, default to ordered for numPr.
        return "ordered"
    return None


def _looks_like_semicolon_item(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if stripped.startswith(("-", "•", ".")):
        return False
    return stripped.endswith(";") or stripped.endswith(".")


def _extract_formula_text(paragraph, plain_text: str) -> str | None:
    # Prefer OMML math content if available.
    math_parts: list[str] = []
    for math_node in paragraph._p.findall(".//m:oMath", paragraph._p.nsmap):
        for t in math_node.findall(".//m:t", paragraph._p.nsmap):
            if t.text:
                math_parts.append(t.text)
    if math_parts:
        return "".join(math_parts).strip() or None

    text = plain_text.strip()
    if _looks_like_formula_text(text):
        return text
    return None


def _looks_like_formula_text(text: str) -> bool:
    if not text or len(text) > 140:
        return False
    if _looks_like_code_text(text):
        return False
    if "=" not in text:
        return False
    # Heuristic markers of math expression.
    return any(marker in text for marker in ("_", "^", "\\", "∑", "π", "√"))


def _looks_like_code_text(text: str) -> bool:
    lowered = text.lower()
    code_markers = (
        "#include",
        "std::",
        "class ",
        "public:",
        "private:",
        "protected:",
        "return ",
        "def ",
        "import ",
        "from ",
        "->",
        "{",
        "}",
    )
    if any(marker in lowered for marker in code_markers):
        return True
    if "(" in text and ")" in text and "=" in text:
        return True
    return False


def _looks_like_caps_heading(text: str) -> bool:
    letters = [ch for ch in text if ch.isalpha()]
    if len(letters) < 3:
        return False
    return text == text.upper()


def _format_term_line(text: str) -> str:
    for sep in (" – ", " — ", " - "):
        if sep in text:
            left, right = text.split(sep, 1)
            left = left.strip()
            right = right.strip()
            if left.lower().startswith("где "):
                symbol = left[4:].strip()
                return f"где {_wrap_var(symbol)}{sep}{right}"
            return f"{_wrap_var(left)}{sep}{right}"
    return text


def _should_join_with_previous(lines: list[str], current: str) -> bool:
    if not lines:
        return False
    if not current:
        return False
    if _looks_like_code_text(current):
        return False
    target_idx = _last_non_empty_index(lines)
    if target_idx is None:
        return False
    prev = lines[target_idx]
    if _looks_like_code_text(prev):
        return False
    if _is_markdown_fence_line(prev):
        return False
    if prev.startswith("#") or prev.startswith("$$"):
        return False
    return _starts_with_lowercase_letter(current)


def _starts_with_lowercase_letter(text: str) -> bool:
    stripped = text.lstrip()
    if not stripped:
        return False
    first = stripped[0]
    return first.isalpha() and first.islower()


def _last_non_empty_index(lines: list[str]) -> int | None:
    for idx in range(len(lines) - 1, -1, -1):
        if lines[idx] != "":
            return idx
    return None


def _wrap_var(symbol: str) -> str:
    s = symbol.strip()
    if not s:
        return s
    if s.startswith("$") and s.endswith("$"):
        return s
    if " " in s:
        return s
    if re.match(r"^[A-Za-zА-Яа-я][A-Za-zА-Яа-я0-9_{}()]*$", s):
        return f"${s}$"
    return s


def _resolve_title_template_path(explicit_path: str | Path | None) -> Path:
    if explicit_path:
        path = Path(explicit_path).expanduser()
        if path.exists():
            return path
        raise FileNotFoundError(f"Title template not found: {path}")
    fallback = Path(__file__).resolve().parents[2] / "examples" / "титульник.docx"
    if fallback.exists():
        return fallback
    raise FileNotFoundError(
        "Title template not found. Expected examples/титульник.docx or explicit title_template_path."
    )


def _find_body_start_index(paragraphs: list) -> int:
    for idx, paragraph in enumerate(paragraphs):
        if _contains_page_break(paragraph):
            return idx + 1
    # If there's no explicit page break, format all content.
    return 0


def _contains_page_break(paragraph) -> bool:
    for run in paragraph.runs:
        for br in run._r.findall(".//w:br", run._r.nsmap):
            if br.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") == "page":
                return True
    return False


def _is_removable_blank(paragraph) -> bool:
    if paragraph.text.strip():
        return False
    # Keep paragraphs that carry page/section structure.
    xml = paragraph._p.xml
    if "w:sectPr" in xml or "w:type=\"page\"" in xml:
        return False
    return True


def _remove_paragraph(paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _extract_paragraph_images(paragraph, images_dir: Path, start_counter: int) -> tuple[list[str], int]:
    refs: list[str] = []
    counter = start_counter
    namespaces = dict(paragraph._p.nsmap)
    namespaces.setdefault("a", "http://schemas.openxmlformats.org/drawingml/2006/main")
    namespaces.setdefault("v", "urn:schemas-microsoft-com:vml")
    rel_ids: list[str] = []
    for blip in paragraph._p.findall(".//a:blip", namespaces):
        rel_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if rel_id:
            rel_ids.append(rel_id)
        rel_link = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link")
        if rel_link:
            rel_ids.append(rel_link)
    for imagedata in paragraph._p.findall(".//v:imagedata", namespaces):
        rel_id = imagedata.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
        if rel_id:
            rel_ids.append(rel_id)
    seen: set[str] = set()
    for rel_id in rel_ids:
        if rel_id in seen:
            continue
        seen.add(rel_id)
        image_part = paragraph.part.related_parts.get(rel_id)
        if image_part is None:
            continue
        suffix = Path(str(getattr(image_part, "partname", ""))).suffix.lower() or ".png"
        file_name = f"img_{counter:04d}{suffix}"
        file_path = images_dir / file_name
        file_path.write_bytes(image_part.blob)
        refs.append(f"images/{file_name}")
        counter += 1
    return refs, counter


def _extract_figure_caption(text: str) -> str | None:
    stripped = text.strip()
    if not stripped:
        return None
    if not stripped.lower().startswith("рисунок"):
        return None
    if "–" in stripped:
        return stripped.split("–", 1)[1].strip() or None
    if "-" in stripped:
        return stripped.split("-", 1)[1].strip() or None
    return None


def _is_markdown_fence_line(text: str) -> bool:
    return re.match(r"^(?:`{3,}|~{3,})\s*$", text.strip()) is not None


def _normalize_markdown_artifacts(text: str) -> str:
    """Drop leaked markdown fence markers from DOCX->MD roundtrips."""
    value = text.strip()
    if not value:
        return ""
    # Remove leaked fence tokens from previous markdown exports.
    value = re.sub(r"(?:`{3,}|~{3,})", "", value).strip()
    return value


def _paragraph_has_image(paragraph) -> bool:
    namespaces = dict(paragraph._p.nsmap)
    namespaces.setdefault("a", "http://schemas.openxmlformats.org/drawingml/2006/main")
    namespaces.setdefault("v", "urn:schemas-microsoft-com:vml")
    if paragraph._p.findall(".//a:blip", namespaces):
        return True
    if paragraph._p.findall(".//v:imagedata", namespaces):
        return True
    return False


def _collect_body_paragraphs_for_extraction(doc: DocxDocument) -> list:
    """Collect paragraphs from body and tables preserving top-level order."""
    body_paragraphs = list(doc.paragraphs)
    body_tables = list(doc.tables)
    p_idx = 0
    t_idx = 0
    ordered: list = []

    for element in doc._body._body.iterchildren():
        tag = element.tag.rsplit("}", 1)[-1]
        if tag == "p":
            if p_idx < len(body_paragraphs):
                ordered.append(body_paragraphs[p_idx])
                p_idx += 1
        elif tag == "tbl":
            if t_idx < len(body_tables):
                table = body_tables[t_idx]
                t_idx += 1
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            ordered.append(paragraph)

    while p_idx < len(body_paragraphs):
        ordered.append(body_paragraphs[p_idx])
        p_idx += 1
    return ordered
