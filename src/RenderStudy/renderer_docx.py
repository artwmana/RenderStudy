from __future__ import annotations

from collections import defaultdict
from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Iterable

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from . import gost_format
from .model import (
    Block,
    CodeBlock,
    Document,
    EquationBlock,
    Heading,
    HorizontalRule,
    ImageBlock,
    InlineEquation,
    InlineElement,
    InlineLink,
    InlineText,
    ListBlock,
    ListItem,
    PageBreak,
    Paragraph,
    TableBlock,
)


@dataclass
class RenderState:
    heading_counters: list[int] = field(default_factory=list)
    current_section: int = 0
    equation_counters: defaultdict[int, int] = field(default_factory=lambda: defaultdict(int))
    figure_counters: defaultdict[int, int] = field(default_factory=lambda: defaultdict(int))
    table_counters: defaultdict[int, int] = field(default_factory=lambda: defaultdict(int))
    asset_root: Path | None = None
    first_heading_rendered: bool = False


LATEX_TO_UNICODE = {
    r"\pi": "π",
    r"\alpha": "α",
    r"\beta": "β",
    r"\gamma": "γ",
    r"\delta": "δ",
    r"\epsilon": "ε",
    r"\theta": "θ",
    r"\lambda": "λ",
    r"\mu": "μ",
    r"\sigma": "σ",
    r"\phi": "φ",
    r"\omega": "ω",
    r"\sum": "∑",
}

SUBSCRIPT_MAP = {
    "0": "₀",
    "1": "₁",
    "2": "₂",
    "3": "₃",
    "4": "₄",
    "5": "₅",
    "6": "₆",
    "7": "₇",
    "8": "₈",
    "9": "₉",
    "a": "ₐ",
    "e": "ₑ",
    "h": "ₕ",
    "i": "ᵢ",
    "j": "ⱼ",
    "k": "ₖ",
    "l": "ₗ",
    "m": "ₘ",
    "n": "ₙ",
    "o": "ₒ",
    "p": "ₚ",
    "r": "ᵣ",
    "s": "ₛ",
    "t": "ₜ",
    "u": "ᵤ",
    "v": "ᵥ",
    "x": "ₓ",
}

SUPERSCRIPT_MAP = {
    "0": "⁰",
    "1": "¹",
    "2": "²",
    "3": "³",
    "4": "⁴",
    "5": "⁵",
    "6": "⁶",
    "7": "⁷",
    "8": "⁸",
    "9": "⁹",
    "+": "⁺",
    "-": "⁻",
    "=": "⁼",
    "(": "⁽",
    ")": "⁾",
    "a": "ᵃ",
    "b": "ᵇ",
    "c": "ᶜ",
    "d": "ᵈ",
    "e": "ᵉ",
    "f": "ᶠ",
    "g": "ᵍ",
    "h": "ʰ",
    "i": "ⁱ",
    "j": "ʲ",
    "k": "ᵏ",
    "l": "ˡ",
    "m": "ᵐ",
    "n": "ⁿ",
    "o": "ᵒ",
    "p": "ᵖ",
    "r": "ʳ",
    "s": "ˢ",
    "t": "ᵗ",
    "u": "ᵘ",
    "v": "ᵛ",
    "w": "ʷ",
    "x": "ˣ",
    "y": "ʸ",
    "z": "ᶻ",
}


def render_document(doc: Document, output_path: str | Path, asset_root: Path | None = None) -> None:
    output_path = Path(output_path)
    state = RenderState(asset_root=asset_root)
    docx = DocxDocument()
    gost_format.apply_page_layout(docx)

    for block in doc.blocks:
        _dispatch_block(docx, block, state)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    docx.save(output_path)


def _dispatch_block(docx: DocxDocument, block: Block, state: RenderState) -> None:
    if isinstance(block, Heading):
        _render_heading(docx, block, state)
    elif isinstance(block, Paragraph):
        _render_paragraph(docx, block.inline)
    elif isinstance(block, ListBlock):
        _render_list(docx, block, state)
    elif isinstance(block, CodeBlock):
        _render_code_block(docx, block)
    elif isinstance(block, EquationBlock):
        _render_equation_block(docx, block, state)
    elif isinstance(block, ImageBlock):
        _render_image_block(docx, block, state)
    elif isinstance(block, HorizontalRule):
        _render_horizontal_rule(docx)
    elif isinstance(block, TableBlock):
        _render_table_block(docx, block, state)
    elif isinstance(block, PageBreak):
        docx.add_page_break()


def _render_heading(docx: DocxDocument, heading: Heading, state: RenderState) -> None:
    if heading.level == 1 and state.first_heading_rendered:
        docx.add_page_break()
    if docx.paragraphs and docx.paragraphs[-1].text.strip():
        spacer_before = docx.add_paragraph("")
        gost_format.apply_body_paragraph_format(spacer_before)
        spacer_before.paragraph_format.first_line_indent = Cm(0)
    number = _compute_heading_number(heading, state)
    heading_text = heading.text.upper() if heading.level == 1 else heading.text
    text = f"{number} {heading_text}" if heading.numbered and number else heading_text
    paragraph = docx.add_paragraph(text)
    style_level = min(max(heading.level, 1), 3)
    paragraph.style = f"Heading {style_level}"
    centered = not heading.numbered
    with_indent = heading.numbered
    for run in paragraph.runs:
        run.bold = True
    gost_format.apply_heading_format(paragraph, centered=centered, with_indent=with_indent)

    # Blank line after heading
    spacer = docx.add_paragraph("")
    gost_format.apply_body_paragraph_format(spacer)
    spacer.paragraph_format.first_line_indent = Cm(0)

    state.first_heading_rendered = True


def _compute_heading_number(heading: Heading, state: RenderState) -> str | None:
    if not heading.numbered:
        return heading.raw_number

    if heading.raw_number:
        try:
            top = int(heading.raw_number.split(".")[0])
            state.current_section = top
        except ValueError:
            pass
        return heading.raw_number

    level = max(1, heading.level)
    while len(state.heading_counters) < level:
        state.heading_counters.append(0)
    state.heading_counters[level - 1] += 1
    for idx in range(level, len(state.heading_counters)):
        state.heading_counters[idx] = 0
    if level == 1:
        state.current_section = state.heading_counters[0]
    number_parts = [str(n) for n in state.heading_counters[:level] if n > 0]
    return ".".join(number_parts)


def _render_paragraph(docx: DocxDocument, inline_elements: Iterable[InlineElement]) -> None:
    paragraph = docx.add_paragraph()
    for inline in inline_elements:
        if isinstance(inline, InlineText):
            run = paragraph.add_run(inline.text)
            gost_format.set_run_font(run, bold=inline.bold, italic=inline.italic, code=inline.code)
        elif isinstance(inline, InlineLink):
            run = paragraph.add_run(inline.text)
            run.font.underline = True
            gost_format.set_run_font(run)
        elif isinstance(inline, InlineEquation):
            run = paragraph.add_run(_latex_to_plain_text(inline.latex))
            gost_format.set_run_font(run)
    gost_format.apply_body_paragraph_format(paragraph)


def _render_list(docx: DocxDocument, block: ListBlock, state: RenderState) -> None:
    # Ensure preceding paragraph ends with colon for list intro
    if docx.paragraphs:
        prev = docx.paragraphs[-1]
        if prev.text and not prev.text.rstrip().endswith(":"):
            prev.text = prev.text.rstrip().rstrip(".,;") + ":"
            for run in prev.runs:
                gost_format.set_run_font(run)

    for idx, item in enumerate(block.items, start=1):
        paragraph = docx.add_paragraph()
        text_parts = []
        if item.blocks and isinstance(item.blocks[0], Paragraph):
            text_parts.append(_inline_to_text(item.blocks[0].inline))
            remaining_blocks = item.blocks[1:]
        else:
            remaining_blocks = item.blocks

        base_text = " ".join(part.strip() for part in text_parts if part.strip())
        formatted_text = _format_list_text(base_text, is_last=idx == len(block.items), ordered=block.ordered)
        prefix = f"{idx} " if block.ordered else "– "
        run = paragraph.add_run(prefix + formatted_text)
        gost_format.set_run_font(run)
        paragraph.paragraph_format.first_line_indent = Cm(gost_format.FIRST_LINE_INDENT_CM)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = Pt(gost_format.LINE_SPACING_PT)
        paragraph.paragraph_format.space_after = Pt(0)

        for sub_block in remaining_blocks:
            _dispatch_block(docx, sub_block, state)


def _render_code_block(docx: DocxDocument, block: CodeBlock) -> None:
    paragraph = docx.add_paragraph()
    run = paragraph.add_run(block.code)
    gost_format.set_run_font(run, code=True)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = Pt(gost_format.LINE_SPACING_PT)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(gost_format.LINE_SPACING_PT)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _render_equation_block(docx: DocxDocument, block: EquationBlock, state: RenderState) -> None:
    # Blank line before
    spacer = docx.add_paragraph("")
    gost_format.apply_body_paragraph_format(spacer)
    spacer.paragraph_format.first_line_indent = Cm(0)

    section = state.current_section or 1
    state.equation_counters[section] += 1
    number = block.number or f"{section}.{state.equation_counters[section]}"

    # Use a 3-column table so the formula is centered on the full text line
    # and the number stays right-aligned at the page margin.
    section = docx.sections[0]
    text_width = section.page_width - section.left_margin - section.right_margin
    text_width_cm = text_width / 360000  # EMU to cm
    number_width_cm = 2.2
    side_width_cm = number_width_cm
    formula_width_cm = max(1.0, text_width_cm - side_width_cm - number_width_cm)

    table = docx.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(side_width_cm)
    table.columns[1].width = Cm(formula_width_cm)
    table.columns[2].width = Cm(number_width_cm)
    _clear_table_borders(table)

    cell_formula = table.cell(0, 1)
    p_formula = cell_formula.paragraphs[0]
    p_formula.paragraph_format.first_line_indent = Cm(0)
    p_formula.paragraph_format.line_spacing = Pt(gost_format.LINE_SPACING_PT)
    p_formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _append_math(p_formula, _latex_to_plain_text(block.latex, convert_scripts=False))

    cell_num = table.cell(0, 2)
    p_num = cell_num.paragraphs[0]
    p_num.paragraph_format.first_line_indent = Cm(0)
    p_num.paragraph_format.line_spacing = Pt(gost_format.LINE_SPACING_PT)
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_number = p_num.add_run(f"({number})")
    gost_format.set_run_font(run_number)

    # Optional description of symbols as aligned list starting with "где"
    if block.terms:
        _render_equation_terms(docx, block.terms)
    # Blank line after
    spacer_after = docx.add_paragraph("")
    gost_format.apply_body_paragraph_format(spacer_after)
    spacer_after.paragraph_format.first_line_indent = Cm(0)


def _render_horizontal_rule(docx: DocxDocument) -> None:
    paragraph = docx.add_paragraph()
    run = paragraph.add_run("-" * 20)
    gost_format.set_run_font(run)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = Pt(gost_format.LINE_SPACING_PT)


def _render_image_block(docx: DocxDocument, block: ImageBlock, state: RenderState) -> None:
    spacer_before = docx.add_paragraph("")
    gost_format.apply_body_paragraph_format(spacer_before)
    spacer_before.paragraph_format.first_line_indent = Cm(0)

    image_path = Path(block.src)
    if state.asset_root:
        candidate = state.asset_root / block.src
        if candidate.exists():
            image_path = candidate

    paragraph = docx.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path))
    except FileNotFoundError:
        run.add_text(f"[Missing image: {image_path}]")
    gost_format.set_run_font(run)

    section = state.current_section or 1
    state.figure_counters[section] += 1
    caption_number = f"{section}.{state.figure_counters[section]}"
    caption_text = block.caption or (block.alt or "").strip() or "Описание рисунка"
    caption_paragraph = docx.add_paragraph(f"Рисунок {caption_number} – {caption_text}")
    gost_format.apply_caption_format(caption_paragraph, centered=True)

    spacer_after = docx.add_paragraph("")
    gost_format.apply_body_paragraph_format(spacer_after)
    spacer_after.paragraph_format.first_line_indent = Cm(0)


def _render_table_block(docx: DocxDocument, block: TableBlock, state: RenderState) -> None:
    section = state.current_section or 1
    state.table_counters[section] += 1
    caption_number = f"{section}.{state.table_counters[section]}"
    caption_text = block.caption or "Название таблицы"
    title = docx.add_paragraph(f"Таблица {caption_number} – {caption_text}")
    gost_format.apply_caption_format(title, space_after=0)

    row_count = len(block.rows)
    col_count = len(block.header) if len(block.header) > 0 else (len(block.rows[0]) if block.rows else 1)
    table = docx.add_table(rows=1 + row_count, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_indent(table, Cm(0.2))
    if block.header:
        for idx, cell_text in enumerate(block.header):
            cell = table.cell(0, idx)
            cell.text = cell_text
    for r_idx, row in enumerate(block.rows, start=1):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = Pt(gost_format.LINE_SPACING_PT)
                for run in paragraph.runs:
                    gost_format.set_run_font(run)

    spacer_after = docx.add_paragraph("")
    gost_format.apply_body_paragraph_format(spacer_after)
    spacer_after.paragraph_format.first_line_indent = Cm(0)


def _latex_to_plain_text(expr: str, convert_scripts: bool = True) -> str:
    """Convert a small subset of LaTeX commands to Unicode glyphs for DOCX text."""
    text = expr.strip()
    for latex, uni in LATEX_TO_UNICODE.items():
        text = text.replace(latex, uni)
    if convert_scripts:
        text = _convert_scripts(text, subscript=True)
        text = _convert_scripts(text, subscript=False)
    return text


def _convert_scripts(text: str, subscript: bool) -> str:
    marker = "_" if subscript else "^"
    mapping = SUBSCRIPT_MAP if subscript else SUPERSCRIPT_MAP
    pattern = rf"\{marker}(\{{[^}}]+\}}|[A-Za-z0-9+\-=()])"

    def repl(match) -> str:
        token = match.group(1)
        value = token[1:-1] if token.startswith("{") and token.endswith("}") else token
        converted = "".join(mapping.get(ch, ch) for ch in value)
        if converted == value:
            return f"{marker}{token}"
        return converted

    return re.sub(pattern, repl, text)


def _inline_to_text(inlines: Iterable[InlineElement]) -> str:
    parts: list[str] = []
    for inline in inlines:
        if isinstance(inline, InlineText):
            parts.append(inline.text)
        elif isinstance(inline, InlineEquation):
            parts.append(inline.latex)
        elif isinstance(inline, InlineLink):
            parts.append(inline.text)
    return "".join(parts)


def _format_list_text(text: str, is_last: bool, ordered: bool) -> str:
    clean = text.strip()
    if not clean:
        return clean
    end = clean[-1]
    if ordered:
        if end != ".":
            clean = clean.rstrip(";") + "."
        return clean
    if end in ".;":
        clean = clean[:-1]
    clean += "." if is_last else ";"
    return clean


def _append_math(paragraph, text: str) -> None:
    """Insert a simple Word math object centered in the paragraph."""
    omath_para = OxmlElement("m:oMathPara")
    omath_para_pr = OxmlElement("m:oMathParaPr")
    jc = OxmlElement("m:jc")
    jc.set(qn("m:val"), "center")
    omath_para_pr.append(jc)
    omath_para.append(omath_para_pr)
    oMath = OxmlElement("m:oMath")
    for node in _build_math_nodes(text):
        oMath.append(node)
    omath_para.append(oMath)
    paragraph._p.append(omath_para)


def _build_math_nodes(text: str) -> list[OxmlElement]:
    nodes: list[OxmlElement] = []
    i = 0
    while i < len(text):
        base_text, i = _read_math_atom(text, i)
        if not base_text:
            break
        sub_text = None
        sup_text = None
        while i < len(text) and text[i] in {"_", "^"}:
            marker = text[i]
            script_text, next_i = _read_math_script(text, i + 1)
            if marker == "_":
                sub_text = script_text
            else:
                sup_text = script_text
            i = next_i

        if sub_text is None and sup_text is None:
            nodes.append(_math_run(base_text))
        elif sub_text is not None and sup_text is not None:
            node = OxmlElement("m:sSubSup")
            e = OxmlElement("m:e")
            e.append(_math_run(base_text))
            node.append(e)
            sub = OxmlElement("m:sub")
            sub.append(_math_run(sub_text))
            node.append(sub)
            sup = OxmlElement("m:sup")
            sup.append(_math_run(sup_text))
            node.append(sup)
            nodes.append(node)
        elif sub_text is not None:
            node = OxmlElement("m:sSub")
            e = OxmlElement("m:e")
            e.append(_math_run(base_text))
            node.append(e)
            sub = OxmlElement("m:sub")
            sub.append(_math_run(sub_text))
            node.append(sub)
            nodes.append(node)
        else:
            node = OxmlElement("m:sSup")
            e = OxmlElement("m:e")
            e.append(_math_run(base_text))
            node.append(e)
            sup = OxmlElement("m:sup")
            sup.append(_math_run(sup_text or ""))
            node.append(sup)
            nodes.append(node)
    return nodes


def _read_math_atom(text: str, start: int) -> tuple[str, int]:
    if start >= len(text):
        return "", start
    if text[start] == "{":
        end = text.find("}", start + 1)
        if end == -1:
            return text[start + 1 :], len(text)
        return text[start + 1 : end], end + 1
    return text[start], start + 1


def _read_math_script(text: str, start: int) -> tuple[str, int]:
    if start >= len(text):
        return "", start
    if text[start] == "{":
        end = text.find("}", start + 1)
        if end == -1:
            return text[start + 1 :], len(text)
        return text[start + 1 : end], end + 1
    return text[start], start + 1


def _math_run(text: str) -> OxmlElement:
    run = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = text
    run.append(t)
    return run


def _render_equation_terms(docx: DocxDocument, terms: list[str]) -> None:
    for idx, term in enumerate(terms):
        symbol, description = _split_term(term)
        description = description.strip()
        if description and not description.endswith(";"):
            description += ";"

        paragraph = docx.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(gost_format.LINE_SPACING_PT)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        prefix_text = "где " if idx == 0 else "       "
        prefix_run = paragraph.add_run(prefix_text)
        gost_format.set_run_font(prefix_run)
        _append_symbol_with_scripts(paragraph, symbol)
        suffix_run = paragraph.add_run(f" –  {description}")
        gost_format.set_run_font(suffix_run)


def _split_term(term: str) -> tuple[str, str]:
    parts = term.split("—", 1)
    if len(parts) == 1:
        parts = term.split("–", 1)
    if len(parts) == 1:
        parts = term.split("-", 1)
    if len(parts) == 2:
        symbol = _normalize_term_symbol(parts[0].strip())
        return symbol, parts[1].strip()
    return _normalize_term_symbol(term.strip()), ""


def _normalize_term_symbol(symbol: str) -> str:
    sym = symbol.strip()
    if sym.startswith("$") and sym.endswith("$") and len(sym) >= 2:
        sym = sym[1:-1].strip()
    return sym


def _append_symbol_with_scripts(paragraph, symbol: str) -> None:
    text = symbol.strip()
    for latex, uni in LATEX_TO_UNICODE.items():
        text = text.replace(latex, uni)

    i = 0
    while i < len(text):
        ch = text[i]
        if ch in {"_", "^"} and i + 1 < len(text):
            is_sub = ch == "_"
            i += 1
            if i < len(text) and text[i] == "{":
                end = text.find("}", i + 1)
                if end == -1:
                    token = text[i + 1 :]
                    i = len(text)
                else:
                    token = text[i + 1 : end]
                    i = end + 1
            else:
                token = text[i]
                i += 1
            run = paragraph.add_run(token)
            gost_format.set_run_font(run)
            run.font.subscript = is_sub
            run.font.superscript = not is_sub
            continue

        start = i
        while i < len(text) and text[i] not in {"_", "^"}:
            i += 1
        run = paragraph.add_run(text[start:i])
        gost_format.set_run_font(run)


def _clear_table_borders(table) -> None:
    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        return
    for child in list(tbl_pr):
        if child.tag == qn("w:tblBorders"):
            tbl_pr.remove(child)
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    tbl_pr.append(borders)


def _set_table_indent(table, indent: Cm) -> None:
    tbl = table._element
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        return
    # remove existing indent
    for child in list(tbl_pr):
        if child.tag == qn("w:tblInd"):
            tbl_pr.remove(child)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), str(int(indent.twips)))
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)
