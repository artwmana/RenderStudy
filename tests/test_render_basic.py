from pathlib import Path

from docx import Document as DocxReader

from RenderStudy.model import Document, EquationBlock, Heading, InlineText, Paragraph
from RenderStudy.renderer_docx import render_document


def test_render_creates_docx(tmp_path: Path):
    doc = Document(
        blocks=[
            Heading(level=1, text="Введение", numbered=True, raw_number="1"),
            Paragraph(inline=[InlineText("Пример абзаца для теста.")]),
            EquationBlock(latex="E = mc^2", display=True),
        ]
    )
    output_file = tmp_path / "report.docx"
    render_document(doc, output_file)
    assert output_file.exists()
    assert output_file.stat().st_size > 0


def test_equation_renders_symbols(tmp_path: Path):
    doc = Document(
        blocks=[
            Heading(level=1, text="Раздел", numbered=True, raw_number="1"),
            EquationBlock(latex="S = \\pi r^2", display=True, terms=["S — площадь круга", "r — радиус"]),
        ]
    )
    out = tmp_path / "eq.docx"
    render_document(doc, out)
    reader = DocxReader(out)
    texts = []
    for p in reader.paragraphs:
        texts.append(p._p.xml)
    # terms are paragraphs after equation
    texts.extend(p._p.xml for p in reader.paragraphs)
    xml = "\n".join(texts)
    assert "π" in xml
    assert "<m:oMath" in xml
    assert "где" in xml
