import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt

from RenderStudy.gost_format import (
    A4_HEIGHT_MM,
    A4_WIDTH_MM,
    FIRST_LINE_INDENT_CM,
    FONT_NAME,
    FONT_SIZE_PT,
    LINE_SPACING,
    MARGIN_BOTTOM_CM,
    MARGIN_LEFT_CM,
    MARGIN_RIGHT_CM,
    MARGIN_TOP_CM,
    apply_body_paragraph_format,
    apply_caption_format,
    apply_heading_format,
    apply_page_layout,
    set_run_font,
)


def test_apply_page_layout():
    doc = Document()
    apply_page_layout(doc)
    section = doc.sections[0]

    # Use EMU (English Metric Unit) tolerance for docx dimensions
    # Cm(1) is 360000 EMUs, Pt(1) is 12700 EMUs.
    # 1000 EMU is a safe tolerance for floating point / rounding differences.
    assert abs(section.page_height - Cm(A4_HEIGHT_MM / 10)) < 1000
    assert abs(section.page_width - Cm(A4_WIDTH_MM / 10)) < 1000
    assert abs(section.left_margin - Cm(MARGIN_LEFT_CM)) < 1000
    assert abs(section.right_margin - Cm(MARGIN_RIGHT_CM)) < 1000
    assert abs(section.top_margin - Cm(MARGIN_TOP_CM)) < 1000
    assert abs(section.bottom_margin - Cm(MARGIN_BOTTOM_CM)) < 1000


def test_apply_body_paragraph_format():
    doc = Document()
    p = doc.add_paragraph("Test body paragraph")
    apply_body_paragraph_format(p)

    assert p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert p.paragraph_format.space_after == Pt(0)
    assert p.paragraph_format.space_before == Pt(0)
    assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
    assert p.paragraph_format.line_spacing == LINE_SPACING
    assert abs(p.paragraph_format.first_line_indent - Cm(FIRST_LINE_INDENT_CM)) < 1000


def test_apply_heading_format():
    doc = Document()
    p = doc.add_paragraph("Test heading")
    apply_heading_format(p, centered=True, with_indent=True)

    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert p.paragraph_format.space_after == Pt(0)
    assert p.paragraph_format.space_before == Pt(0)
    assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
    assert p.paragraph_format.page_break_before is False
    assert abs(p.paragraph_format.left_indent - Cm(0)) < 1000
    assert abs(p.paragraph_format.right_indent - Cm(0)) < 1000
    assert abs(p.paragraph_format.first_line_indent - Cm(FIRST_LINE_INDENT_CM)) < 1000

    for run in p.runs:
        assert run.font.name == FONT_NAME
        assert run.font.size == Pt(FONT_SIZE_PT)
        assert run.bold is True
        assert run.italic is False


def test_apply_heading_format_not_centered():
    doc = Document()
    p = doc.add_paragraph("Test heading")
    apply_heading_format(p, centered=False, with_indent=False)

    assert p.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert p.paragraph_format.first_line_indent is None or abs(p.paragraph_format.first_line_indent - Cm(0)) < 1000


def test_apply_caption_format():
    doc = Document()
    p = doc.add_paragraph("Test caption")
    apply_caption_format(p, centered=True, space_after=10)

    assert p.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert p.paragraph_format.space_after == Pt(10)
    assert p.paragraph_format.space_before == Pt(0)
    assert p.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
    assert p.paragraph_format.first_line_indent is None or abs(p.paragraph_format.first_line_indent - Cm(0)) < 1000

    for run in p.runs:
        assert run.font.name == FONT_NAME
        assert run.font.size == Pt(FONT_SIZE_PT)
        assert run.bold is False
        assert run.italic is False


def test_set_run_font():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Test text")
    set_run_font(run, bold=True, italic=True)

    assert run.font.name == FONT_NAME
    assert run.font.size == Pt(FONT_SIZE_PT)
    assert run.bold is True
    assert run.italic is True
