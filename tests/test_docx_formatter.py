from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt

from RenderStudy.docx_formatter import rebuild_docx_via_markdown, reformat_docx


def test_reformat_docx_applies_font_and_spacing(tmp_path: Path):
    src = tmp_path / "input.docx"
    out = tmp_path / "output.docx"

    doc = DocxDocument()
    p1 = doc.add_paragraph("ВВЕДЕНИЕ")
    p1.runs[0].font.name = "Calibri"
    p1.runs[0].font.size = Pt(11)
    p2 = doc.add_paragraph("Обычный абзац текста.")
    p2.runs[0].font.name = "Calibri"
    p2.runs[0].font.size = Pt(11)
    doc.save(src)

    reformat_docx(src, out)

    formatted = DocxDocument(out)
    heading = formatted.paragraphs[0]
    body = formatted.paragraphs[1]

    assert heading.runs[0].font.name == "Times New Roman"
    assert heading.runs[0].font.size.pt == 14.0
    assert body.runs[0].font.name == "Times New Roman"
    assert body.runs[0].font.size.pt == 14.0
    assert heading.paragraph_format.line_spacing == 1.0
    assert body.paragraph_format.line_spacing == 1.0


def test_reformat_docx_keeps_title_page_and_cleans_body_blanks(tmp_path: Path):
    src = tmp_path / "input_with_title.docx"
    out = tmp_path / "output_with_title.docx"

    doc = DocxDocument()
    title = doc.add_paragraph("ТИТУЛЬНИК")
    title.runs[0].font.name = "Calibri"
    title.runs[0].font.size = Pt(11)
    doc.add_page_break()
    doc.add_paragraph("")
    body = doc.add_paragraph("1 ВВЕДЕНИЕ")
    body.runs[0].font.name = "Calibri"
    body.runs[0].font.size = Pt(11)
    doc.save(src)

    reformat_docx(src, out)

    formatted = DocxDocument(out)
    texts = [p.text for p in formatted.paragraphs]
    assert "ТИТУЛЬНИК" in texts
    assert "1 ВВЕДЕНИЕ" in texts
    # Only one blank paragraph remains: service page-break marker.
    blank_paragraphs = [p for p in formatted.paragraphs if p.text == ""]
    assert len(blank_paragraphs) == 1
    assert 'w:type="page"' in blank_paragraphs[0]._p.xml


def test_rebuild_docx_via_markdown_uses_external_title_template(tmp_path: Path):
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    md_dump = tmp_path / "body.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("ШАБЛОННЫЙ ТИТУЛ")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("1 ВВЕДЕНИЕ")
    doc.add_paragraph("Обычный текст раздела.")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)

    assert out.exists()
    assert md_dump.exists()
    md_text = md_dump.read_text(encoding="utf-8")
    assert "# 1 ВВЕДЕНИЕ" in md_text
    assert "Обычный текст раздела." in md_text

    rebuilt = DocxDocument(out)
    texts = [p.text for p in rebuilt.paragraphs]
    assert "ШАБЛОННЫЙ ТИТУЛ" in texts
    assert "СТАРЫЙ ТИТУЛ" not in texts
    assert "1 ВВЕДЕНИЕ" in texts


def test_rebuild_docx_via_markdown_trims_old_title_tail_and_keeps_headings(tmp_path: Path):
    src = tmp_path / "in_tail.docx"
    out = tmp_path / "out_tail.docx"
    md_dump = tmp_path / "tail.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("НОВЫЙ ТИТУЛ")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("Белорусский государственный университет ...")
    doc.add_paragraph("Кафедра ...")
    p_h = doc.add_paragraph("1 ВВЕДЕНИЕ")
    p_h.style = "Heading 1"
    doc.add_paragraph("Цель работы — разработка...")
    doc.add_paragraph("Текст раздела.")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "Белорусский государственный" not in md_text
    assert "# 1 ВВЕДЕНИЕ" in md_text
    assert "Цель работы" in md_text


def test_rebuild_docx_via_markdown_formats_caps_heading_and_terms(tmp_path: Path):
    src = tmp_path / "in_caps.docx"
    out = tmp_path / "out_caps.docx"
    md_dump = tmp_path / "caps.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("ШАБЛОН")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("где S_ПП – площадь печатной платы;")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "# ВВЕДЕНИЕ" in md_text
    assert "где $S_ПП$ – площадь печатной платы;" in md_text


def test_rebuild_docx_via_markdown_no_page_break_drops_old_title(tmp_path: Path):
    src = tmp_path / "in_no_break.docx"
    out = tmp_path / "out_no_break.docx"
    md_dump = tmp_path / "no_break.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("НОВЫЙ ТИТУЛ")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("МИНИСТЕРСТВО ...")
    doc.add_paragraph("БГУИР")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст раздела.")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "МИНИСТЕРСТВО" not in md_text
    assert "БГУИР" not in md_text
    assert "# ВВЕДЕНИЕ" in md_text


def test_rebuild_docx_via_markdown_joins_lowercase_lines(tmp_path: Path):
    src = tmp_path / "in_join.docx"
    out = tmp_path / "out_join.docx"
    md_dump = tmp_path / "join.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("ШАБЛОН")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("Создать класс Sphere (шар).")
    doc.add_paragraph("поле класса хранит радиус шара.")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "Создать класс Sphere (шар). поле класса хранит радиус шара." in md_text
    assert "Создать класс Sphere (шар).\n\nполе класса хранит радиус шара." not in md_text


def test_rebuild_docx_via_markdown_code_line_not_formula(tmp_path: Path):
    src = tmp_path / "in_code.docx"
    out = tmp_path / "out_code.docx"
    md_dump = tmp_path / "code.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("ШАБЛОН")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("return radius * radius;")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "$$return radius * radius;$$" not in md_text
    assert "return radius * radius;" in md_text


def test_rebuild_docx_via_markdown_code_line_not_joined(tmp_path: Path):
    src = tmp_path / "in_code_join.docx"
    out = tmp_path / "out_code_join.docx"
    md_dump = tmp_path / "code_join.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("ШАБЛОН")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("Описание алгоритма.")
    doc.add_paragraph("return radius * radius;")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "Описание алгоритма. return radius * radius;" not in md_text
    assert "Описание алгоритма.\n\n```\nreturn radius * radius;\n```" in md_text


def test_rebuild_docx_via_markdown_keeps_numbered_lists(tmp_path: Path):
    src = tmp_path / "in_list.docx"
    out = tmp_path / "out_list.docx"
    md_dump = tmp_path / "list.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("ШАБЛОН")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("Цель работы:")
    p1 = doc.add_paragraph("Первый пункт")
    p1.style = "List Number"
    p2 = doc.add_paragraph("Второй пункт")
    p2.style = "List Number"
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "1. Первый пункт" in md_text
    assert "1. Второй пункт" in md_text


def test_rebuild_docx_via_markdown_colon_semicolon_to_bullets(tmp_path: Path):
    src = tmp_path / "in_semicolon_list.docx"
    out = tmp_path / "out_semicolon_list.docx"
    md_dump = tmp_path / "semicolon_list.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("ШАБЛОН")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("Основные параметры:")
    doc.add_paragraph("напряжение питания 3,7 В;")
    doc.add_paragraph("рабочая температура 0..40 С;")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "Основные параметры:" in md_text
    assert "- напряжение питания 3,7 В;" in md_text
    assert "- рабочая температура 0..40 С;" in md_text


def test_rebuild_docx_via_markdown_colon_semicolon_final_dot_stays_same_list(tmp_path: Path):
    src = tmp_path / "in_semicolon_dot_list.docx"
    out = tmp_path / "out_semicolon_dot_list.docx"
    md_dump = tmp_path / "semicolon_dot_list.md"
    template = tmp_path / "title_template.docx"

    tpl = DocxDocument()
    tpl.add_paragraph("ШАБЛОН")
    tpl.save(template)

    doc = DocxDocument()
    doc.add_paragraph("СТАРЫЙ ТИТУЛ")
    doc.add_page_break()
    doc.add_paragraph("Система включает:")
    p1 = doc.add_paragraph("первый элемент;")
    p1.style = "List Number"
    p2 = doc.add_paragraph("второй элемент.")
    p2.style = "List Number"
    doc.add_paragraph("Дальше идет обычный текст.")
    doc.save(src)

    rebuild_docx_via_markdown(src, out, extracted_md_path=md_dump, title_template_path=template)
    md_text = md_dump.read_text(encoding="utf-8")

    assert "- первый элемент;" in md_text
    assert "- второй элемент." in md_text
    assert "1. второй элемент." not in md_text
