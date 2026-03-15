import time
from docx import Document
from RenderStudy.docx_formatter import reformat_docx
import tempfile
import os

def generate_large_docx(path, num_paragraphs):
    doc = Document()
    for i in range(num_paragraphs):
        doc.add_paragraph(f"Paragraph {i}")
        if i % 2 == 0:
            doc.add_paragraph("")
    doc.save(path)

def benchmark():
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, "large.docx")
        output_path = os.path.join(tmpdir, "out.docx")

        generate_large_docx(input_path, 10000)

        start_time = time.time()
        reformat_docx(input_path, output_path)
        end_time = time.time()

        print(f"Time taken: {end_time - start_time:.4f} seconds")

if __name__ == "__main__":
    benchmark()
